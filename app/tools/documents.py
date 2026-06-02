"""Document tools deterministas (T5) — generan DOCX profesional + markdown.

El modelo decide el CONTENIDO; estas funciones producen el ARCHIVO (formato
garantizado). Header de work-product inyectado en todos.
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt

WP_HEADER = "PRIVILEGED & CONFIDENTIAL — ATTORNEY WORK PRODUCT"


def _new_doc() -> Document:
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run(WP_HEADER)
    r.bold = True
    r.font.size = Pt(9)
    return doc


def _bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_memo(title: str, sections: list[dict], **_) -> tuple[bytes, str]:
    doc = _new_doc()
    doc.add_heading(title, level=0)
    md = [f"**{WP_HEADER}**", "", f"# {title}", ""]
    for s in sections or []:
        h, b = s.get("heading"), s.get("body", "")
        if h:
            doc.add_heading(h, level=1)
            md.append(f"## {h}")
        for para in (b or "").split("\n"):
            doc.add_paragraph(para)
        md += [b or "", ""]
    return _bytes(doc), "\n".join(md)


def render_letter(recipient: str, body: str, sender: str = "", date: str = "", **_) -> tuple[bytes, str]:
    doc = _new_doc()
    md = [f"**{WP_HEADER}**", ""]
    if date:
        doc.add_paragraph(date)
        md.append(date)
    doc.add_paragraph(f"Para: {recipient}")
    md.append(f"Para: {recipient}")
    doc.add_paragraph("")
    for para in (body or "").split("\n"):
        doc.add_paragraph(para)
    md += ["", body or ""]
    if sender:
        doc.add_paragraph("")
        doc.add_paragraph(sender)
        md += ["", sender]
    return _bytes(doc), "\n".join(md)


def build_table_doc(title: str, columns: list[str], rows: list[list], **_) -> tuple[bytes, str]:
    doc = _new_doc()
    doc.add_heading(title, level=0)
    cols = columns or []
    table = doc.add_table(rows=1, cols=max(1, len(cols)))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = str(c)
    md = [f"**{WP_HEADER}**", "", f"# {title}", "", "| " + " | ".join(map(str, cols)) + " |",
          "|" + "---|" * len(cols)]
    for row in rows or []:
        cells = table.add_row().cells
        for i in range(len(cols)):
            cells[i].text = str(row[i]) if i < len(row) else ""
        md.append("| " + " | ".join(str(row[i]) if i < len(row) else "" for i in range(len(cols))) + " |")
    return _bytes(doc), "\n".join(md)


GENERATORS = {
    "render_memo": render_memo,
    "render_letter": render_letter,
    "build_table_doc": build_table_doc,
}
